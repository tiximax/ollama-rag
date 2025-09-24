import os
import io
import tempfile
from typing import List, Dict, Any

import pytest

from app.rag_engine import (
    chunk_text,
    extract_text_from_docx,
    extract_text_from_pdf,
    RagEngine,
)


def test_chunk_text_basic_overlap():
    text = "abcdefghij"  # len=10
    chunks = chunk_text(text, chunk_size=4, overlap=2)
    # Expected windows: [0:4] abcd, then start=4-2=2 -> [2:6] cdef, then [4:8] efgh, then [6:10] ghij
    assert chunks == ["abcd", "cdef", "efgh", "ghij"]


def test_chunk_text_exact_fit():
    text = "abcdefgh"  # len=8
    chunks = chunk_text(text, chunk_size=4, overlap=2)
    assert chunks == ["abcd", "cdef", "efgh"]


def test_extract_docx_success(tmp_path):
    # Create a small .docx file programmatically
    try:
        from docx import Document  # type: ignore
    except Exception:
        pytest.skip("python-docx is not available")
    d = Document()
    d.add_paragraph("Xin chao")
    d.add_paragraph("RAG test")
    path = tmp_path / "sample.docx"
    d.save(str(path))
    out = extract_text_from_docx(str(path))
    # Lines joined by \n
    assert "Xin chao" in out
    assert "RAG test" in out


def test_extract_pdf_fallback_returns_empty(tmp_path):
    # Write a non-PDF file with .pdf extension; extractor should fail and return ""
    p = tmp_path / "not_a_pdf.pdf"
    p.write_text("this is not really a pdf")
    out = extract_text_from_pdf(str(p))
    assert out == ""


class DummyEngine(RagEngine):
    """Override retrieval to avoid Chroma/Ollama, used only for logic tests."""

    def __init__(self):
        # Do not call parent __init__ to avoid Chroma, just set minimal fields
        self.collection_name = "docs"
        self.default_provider = "ollama"
        self._bge_rr = None
        self._embed_rr = None
        # minimal stores
        self.persist_root = "."
        self.db_name = "default"

    def retrieve(self, query: str, top_k: int = 5, *, languages=None, versions=None) -> Dict[str, Any]:
        # Return deterministic vector result: v_docs first
        docs = [f"V-{query}-A", f"V-{query}-B", f"V-{query}-C"][:top_k]
        metas = [{"source": f"s-{i}", "chunk": i} for i, _ in enumerate(docs)]
        distances = [0.1, 0.2, 0.3][: len(docs)]
        return {"documents": docs, "metadatas": metas, "distances": distances}

    def retrieve_bm25(self, query: str, top_k: int = 5, *, languages=None, versions=None) -> Dict[str, Any]:
        docs = [f"B-{query}-X", f"B-{query}-Y", f"B-{query}-Z"][:top_k]
        metas = [{"source": f"b-{i}", "chunk": i} for i, _ in enumerate(docs)]
        scores = [3.0, 2.0, 1.0][: len(docs)]
        return {"documents": docs, "metadatas": metas, "scores": scores}

    def retrieve_hybrid(self, query: str, top_k: int = 5, bm25_weight: float = 0.5, rrf_enable=None, rrf_k=None, *, languages=None, versions=None) -> Dict[str, Any]:
        # Simple mix: interleave V- and B- prefixes
        v = self.retrieve(query, top_k=top_k)
        b = self.retrieve_bm25(query, top_k=top_k)
        docs = []
        metas = []
        for i in range(top_k):
            if i < len(v["documents"]):
                docs.append(v["documents"][i])
                metas.append(v["metadatas"][i])
            if i < len(b["documents"]):
                docs.append(b["documents"][i])
                metas.append(b["metadatas"][i])
        return {"documents": docs[:top_k], "metadatas": metas[:top_k]}

    def _get_llm(self, provider=None):  # avoid real LLM
        class _LLM:
            def generate(self, prompt):
                return "ok"

            def generate_stream(self, prompt):
                yield "ok"

        return _LLM()


def test_rerank_fallback_embed_monkeypatch(monkeypatch):
    eng = DummyEngine()

    # Monkeypatch embedder to produce similarity that ranks doc2 highest
    calls = {"count": 0}

    def fake_embed(texts: List[str]) -> List[List[float]]:
        # first is query, then docs; craft cosine similarities:
        # q=[1,0], d1=[0,1], d2=[1,0] -> cosine(q,d2)=1.0 > cosine(q,d1)=0.0
        out = []
        for idx, _ in enumerate(texts):
            if idx == 0:  # query
                out.append([1.0, 0.0])
            elif idx == 1:  # doc1
                out.append([0.0, 1.0])
            else:  # doc2...
                out.append([1.0, 0.0])
        calls["count"] += 1
        return out

    # inject SimpleEmbedReranker with our fake embed
    from app.rag_engine import SimpleEmbedReranker

    eng._embed_rr = SimpleEmbedReranker(fake_embed)

    docs = ["d1", "d2"]
    metas = [{"source": "s1", "chunk": 0}, {"source": "s2", "chunk": 1}]
    ranked_docs, ranked_metas = eng._apply_rerank("q", docs, metas, top_k=2, rr_provider="embed")
    assert ranked_docs[0] == "d2"  # due to cosine similarity
    assert len(ranked_docs) == 2


def test_retrieve_aggregate_rrf_across_rewrites(monkeypatch):
    eng = DummyEngine()

    # Force rewrites to two variants
    monkeypatch.setattr(eng, "_rewrite_queries", lambda q, n=2, provider=None: ["q1", "q2"])

    # For q1, prefer vector docs; for q2, also return vector docs different so RRF unions
    def ret_vec(q, top_k=5, languages=None, versions=None):
        docs = [f"V-{q}-1", f"V-{q}-2"]
        metas = [{"source": f"{q}-s1", "chunk": 0}, {"source": f"{q}-s2", "chunk": 1}]
        # distances smaller => higher similarity
        return {"documents": docs, "metadatas": metas, "distances": [0.1, 0.2]}

    monkeypatch.setattr(eng, "retrieve", ret_vec)

    res = eng.retrieve_aggregate("q0", top_k=3, method="vector", rewrite_enable=True, rewrite_n=2)
    docs = res.get("documents", [])
    metas = res.get("metadatas", [])
    assert len(docs) <= 3
    # Expect unique union (some from q1, some from q2)
    assert any("V-q1-" in d for d in docs)
    assert any("V-q2-" in d for d in docs)


def test_valid_db_name():
    # Use staticmethod via class
    assert RagEngine._valid_db_name("abc-123_45")
    assert not RagEngine._valid_db_name("")
    assert not RagEngine._valid_db_name("../../etc/passwd")
    assert not RagEngine._valid_db_name("bad name with spaces")
    assert not RagEngine._valid_db_name("a" * 65)  # over 64 chars
